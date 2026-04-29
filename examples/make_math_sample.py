"""
生成带 OMML 公式的 docx——模拟学生用 Word 公式编辑器画的公式。
python-docx 不直接支持构造 OMML，但我们可以直接往 XML 里塞 OMML 片段。
"""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from lxml import etree

OUT = Path(__file__).parent / "math_sample.docx"

doc = Document()
doc.add_heading("含公式的样本", level=1)
doc.add_paragraph("以下段落含 Word 公式编辑器画的公式：")

# 在段落里嵌入一个 OMML 块（行内公式：a^2 + b^2 = c^2）
p = doc.add_paragraph("勾股定理 ")

# OMML 命名空间
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
nsmap_m = {"m": M_NS}

# <m:oMath> a^2 + b^2 = c^2 </m:oMath>
omml_inline = etree.SubElement(p._p, qn("m:oMath"))
def m(tag): return etree.SubElement_(parent=None, tag=qn(f"m:{tag}"))  # placeholder; we'll build manually

# 重写：手工构建 OMML 子树
def build_sup(parent, base_text, exp_text):
    """构造 a^2 这种上标。"""
    sSup = etree.SubElement(parent, qn("m:sSup"))
    e = etree.SubElement(sSup, qn("m:e"))
    r1 = etree.SubElement(e, qn("m:r"))
    t1 = etree.SubElement(r1, qn("m:t")); t1.text = base_text
    sup = etree.SubElement(sSup, qn("m:sup"))
    r2 = etree.SubElement(sup, qn("m:r"))
    t2 = etree.SubElement(r2, qn("m:t")); t2.text = exp_text

def add_run_text(parent, text):
    r = etree.SubElement(parent, qn("m:r"))
    t = etree.SubElement(r, qn("m:t")); t.text = text

# 清空 oMath 重新构造
omml_inline.clear()
build_sup(omml_inline, "a", "2")
add_run_text(omml_inline, " + ")
build_sup(omml_inline, "b", "2")
add_run_text(omml_inline, " = ")
build_sup(omml_inline, "c", "2")

p.add_run(" 是几何中的著名定理。")

# 第二段：含一个块级 OMML（求和公式）
p2 = doc.add_paragraph("求和公式（块级）：")
oMathPara = etree.SubElement(p2._p, qn("m:oMathPara"))
oMath = etree.SubElement(oMathPara, qn("m:oMath"))

# Σ_{i=1}^{n} i = n(n+1)/2
nary = etree.SubElement(oMath, qn("m:nary"))
naryPr = etree.SubElement(nary, qn("m:naryPr"))
chr_el = etree.SubElement(naryPr, qn("m:chr")); chr_el.set(qn("m:val"), "∑")
sub = etree.SubElement(nary, qn("m:sub"))
add_run_text(sub, "i=1")
sup = etree.SubElement(nary, qn("m:sup"))
add_run_text(sup, "n")
e = etree.SubElement(nary, qn("m:e"))
add_run_text(e, "i")
add_run_text(oMath, " = ")
# 分数 n(n+1)/2
f_el = etree.SubElement(oMath, qn("m:f"))
num = etree.SubElement(f_el, qn("m:num"))
add_run_text(num, "n(n+1)")
den = etree.SubElement(f_el, qn("m:den"))
add_run_text(den, "2")

doc.add_paragraph("公式之后还有一段普通文字。")
doc.add_paragraph("学生也可以用文本写：$E = m c^2$ 这种简单形式。")

doc.save(OUT)
print(f"Sample written: {OUT}")

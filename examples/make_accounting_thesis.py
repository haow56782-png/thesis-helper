"""
会计学专业本科毕业论文样本生成器
课题：数字经济背景下中小企业财务共享服务中心建设研究——以 X 公司为例
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parent / "accounting_thesis.docx"

doc = Document()
style = doc.styles["Normal"]
style.font.name = "宋体"
style.font.size = Pt(12)


def heading(text, level=1):
    doc.add_heading(text, level=level)


def para(text, indent=True):
    p = doc.add_paragraph(text)
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    return p


def pseudo_title(text, size=16):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_table(headers, rows, style_name="Table Grid"):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = style_name
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    for i, row in enumerate(rows, start=1):
        for j, cell_text in enumerate(row):
            table.rows[i].cells[j].text = str(cell_text)
    return table


# ============ 中文摘要 ============
pseudo_title("摘 要")
para("""随着数字经济的快速发展和企业管理模式的不断升级，财务共享服务中心（Financial Shared Service Center, FSSC）作为一种集中化、标准化、信息化的财务管理新模式，正在被越来越多的企业所采纳。中小企业在数字化转型过程中面临成本控制、流程优化、人才短缺等多重挑战，建设适合自身规模的财务共享服务中心成为破解这些难题的有效路径。""")

para("""本文以 X 公司为研究对象，系统梳理了财务共享服务中心的理论基础与实践模式，深入分析了 X 公司在传统财务管理模式下存在的问题。针对存在的问题，本文设计了适合 X 公司的财务共享服务中心建设方案，包括组织架构、业务流程、信息系统和绩效评价体系四个方面。通过对实施效果的评估，本文发现财务共享服务中心建设有效降低了 X 公司的财务成本，提升了财务工作效率，强化了内部控制水平。""")

para("""本研究为同类型中小企业建设财务共享服务中心提供了参考，对推动企业财务数字化转型具有一定的实践意义。""")

p = doc.add_paragraph()
r = p.add_run("关键词：")
r.bold = True
p.add_run("数字经济；财务共享服务中心；中小企业；财务管理；数字化转型")

# ============ 英文摘要 ============
pseudo_title("Abstract")
para("""With the rapid development of the digital economy and the continuous upgrading of enterprise management models, the Financial Shared Service Center (FSSC), as a centralized, standardized, and information-based new model of financial management, is being adopted by more and more enterprises. Small and medium-sized enterprises (SMEs) face multiple challenges such as cost control, process optimization, and talent shortage in the process of digital transformation, and building an FSSC suitable for their own scale has become an effective path to solve these problems.""")

para("""This paper takes Company X as the research object, systematically reviews the theoretical foundation and practical models of FSSC, and deeply analyzes the problems existing in Company X under the traditional financial management mode. In view of the existing problems, this paper designs an FSSC construction plan suitable for Company X, including four aspects: organizational structure, business process, information system, and performance evaluation system.""")

para("""This study provides reference for SMEs of similar types to build FSSCs, and has certain practical significance for promoting the digital transformation of enterprise finance.""")

p = doc.add_paragraph()
r = p.add_run("Keywords: ")
r.bold = True
p.add_run("Digital Economy; Financial Shared Service Center; SMEs; Financial Management; Digital Transformation")


# ============ 第一章 绪论 ============
heading("绪论", level=1)

heading("研究背景与意义", level=2)
para("""进入 21 世纪以来，全球经济正在经历以数字化、网络化、智能化为特征的深刻变革。国家"十四五"规划明确提出要加快数字经济发展，推动数字技术与实体经济深度融合 [1]。在此背景下，企业的财务管理模式也面临着前所未有的转型压力。传统的分散式财务管理模式由于流程冗长、信息孤岛、成本高昂等弊端，已经难以适应数字经济时代企业对高效、透明、敏捷财务管理的需求 [2]。""")

para("""财务共享服务中心作为现代企业财务管理的创新模式，最早起源于 20 世纪 80 年代的福特公司。经过四十多年的发展，FSSC 已经在大型跨国公司中得到广泛应用，并取得了显著成效 [3]。然而，对于数量庞大的中小企业而言，由于资源有限、规模较小，如何借鉴 FSSC 的核心理念建设适合自身的财务管理体系，仍然是一个值得深入研究的课题。""")

para("""本研究的理论意义在于：第一，丰富了中小企业财务共享服务中心建设的理论体系；第二，将数字经济与财务管理深度融合，为相关领域的交叉研究提供了新的视角。本研究的实践意义在于：通过对 X 公司的案例分析，为同类型中小企业建设 FSSC 提供可借鉴的经验和参考方案。""")

heading("国内外研究现状", level=2)
para("""国外学者对财务共享服务中心的研究起步较早。Quinn 等 [4] 系统阐述了共享服务的概念框架，认为 FSSC 的本质是通过流程标准化和服务集中化降低成本、提升效率。Ulbrich [5] 通过对欧美 200 家企业的实证研究发现，成功的 FSSC 实施能够使企业财务成本降低 25% 至 40%。近年来，国外研究开始关注 FSSC 与新兴技术的融合，如机器人流程自动化（RPA）、人工智能、区块链等技术对 FSSC 的赋能作用 [6]。""")

para("""国内研究虽然起步较晚，但发展迅速。张瑞君 [7] 较早将 FSSC 概念引入国内，结合中国国有企业改革背景探讨了 FSSC 的本土化路径。陈虎、孙彦丛 [8] 从信息技术角度系统研究了 FSSC 的信息系统建设问题，提出了业务流程驱动加信息系统支撑加共享服务实施的三层架构模型。近年来，针对中小企业 FSSC 建设的研究逐渐增多 [9,10]，但仍存在理论研究与实践脱节、可操作性不足等问题 [11-12]。""")

heading("研究内容与方法", level=2)
para("""本文以"数字经济背景下中小企业财务共享服务中心建设"为研究主题，以 X 公司为案例对象，采用以下研究方法：""")

doc.add_paragraph("文献研究法：通过查阅国内外相关文献，梳理 FSSC 的理论发展脉络和最新研究成果；", style="List Number")
doc.add_paragraph("案例研究法：以 X 公司为典型案例，深入剖析其财务管理现状和共享服务中心建设方案；", style="List Number")
doc.add_paragraph("比较分析法：通过对比 FSSC 实施前后的关键财务指标，评估实施效果；", style="List Number")
doc.add_paragraph("调查访谈法：通过对 X 公司财务部门相关人员的访谈，获取一手研究资料。", style="List Number")


# ============ 第二章 理论基础 ============
heading("财务共享服务中心理论基础", level=1)

heading("财务共享服务中心的概念界定", level=2)
para("""财务共享服务中心是指企业将分散在各个业务单元中的相同或相似的财务业务，通过流程再造与信息系统支持的方式集中到一个统一的服务平台进行处理的组织模式。其核心理念可以概括为集中化、标准化、专业化、信息化四个维度 [3]。""")

para("""FSSC 与传统财务部门的本质区别在于服务模式的转变。传统财务部门面向企业内部各业务单元提供行政管理式服务，而 FSSC 则采用客户服务式模式，将业务单元视为内部客户，通过签订服务水平协议（Service Level Agreement, SLA）的方式约定服务标准 [4]。""")

heading("财务共享服务中心的运作模式", level=2)
para("""根据集中化程度和服务范围的不同，财务共享服务中心主要有以下三种运作模式：""")
para("""第一，基础型 FSSC。主要集中处理应付账款、应收账款、费用报销、总账核算等基础性、重复性的财务业务，约占总财务业务量的 60% 至 70%。该模式适合处于 FSSC 建设初期的企业。""")
para("""第二，扩展型 FSSC。在基础型的基础上，进一步将税务管理、资金管理、固定资产管理等业务纳入共享范围，并开始向业务单元提供决策支持性服务。""")
para("""第三，战略型 FSSC。集成预算管理、绩效管理、风险管理等高端财务业务，向企业管理层提供战略决策支持，是 FSSC 发展的高级阶段。""")

heading("数字经济对财务共享的赋能", level=2)
para("""数字经济时代涌现的一系列新兴技术，正在深刻改变 FSSC 的运作方式 [6]：""")
add_table(
    headers=["技术类型", "应用场景", "赋能效果"],
    rows=[
        ["机器人流程自动化（RPA）", "凭证录入、报表合并", "效率提升 60% 以上"],
        ["人工智能（AI）", "票据识别、异常检测", "准确率达 99% 以上"],
        ["大数据分析", "财务预测、风险预警", "决策响应速度提升 3 倍"],
        ["区块链技术", "电子凭证、跨境结算", "可信度与透明度显著增强"],
        ["云计算", "系统部署、数据存储", "IT 成本降低 30% 至 50%"],
    ],
)


# ============ 第三章 X 公司现状分析 ============
heading("X 公司财务管理现状分析", level=1)

heading("X 公司基本情况", level=2)
para("""X 公司成立于 2010 年，是一家从事高端制造业的中小型民营企业，总部位于重庆市，在成都、武汉、苏州设有 3 家分公司，员工总数约 800 人。公司近三年的主要财务指标如下表所示：""")

add_table(
    headers=["指标", "2022 年", "2023 年", "2024 年"],
    rows=[
        ["营业收入（万元）", "12,460", "15,820", "18,950"],
        ["净利润（万元）", "1,120", "1,560", "2,030"],
        ["资产总额（万元）", "23,580", "28,940", "34,210"],
        ["员工人数", "640", "720", "800"],
        ["财务人员数", "18", "22", "26"],
    ],
)

heading("现有财务管理模式分析", level=2)
para("""X 公司目前采用的是分散式财务管理模式。总部设有财务部，各分公司设有独立的财务科室。这种模式在公司规模较小时能够保证业务的灵活性，但随着公司规模的扩大和业务的复杂化，逐渐暴露出诸多问题。""")

heading("存在的主要问题", level=2)
para("""通过对 X 公司财务部门的深入调研和访谈，本文识别出以下四个主要问题：""")
para("""第一，财务流程不统一。各分公司在费用报销、应付管理等核心流程上存在差异，总部在合并报表和数据汇总时需要进行大量的口径调整，工作量大且容易出错。""")
para("""第二，财务成本居高不下。当前财务人员占员工总数的 3.25%，远高于行业平均水平 1.8%。财务工作中的重复性劳动占比高达 70%，导致人力成本浪费严重。""")
para("""第三，信息系统割裂。各分公司使用的财务软件版本不统一，数据交换需要人工导出导入，财务数据的及时性和准确性难以保证 [10]。""")
para("""第四，内部控制薄弱。分散式管理使得职能交叉、权责不清的现象时有发生，近三年累计发生差错事项 23 起，造成直接经济损失约 56 万元。""")


# ============ 第四章 建设方案设计 ============
heading("X 公司财务共享服务中心建设方案", level=1)

heading("建设目标与原则", level=2)
para("""结合 X 公司的实际情况，本文提出 FSSC 建设的总体目标是：通过流程再造、组织调整和信息系统支持，在三年内将财务工作效率提升 50% 以上，财务成本降低 30% 以上，差错率控制在 0.5% 以内，建立起符合公司发展战略的现代财务管理体系。""")

para("""建设过程应遵循统一规划、分步实施、效益优先、风险可控四大原则。""")

heading("组织架构设计", level=2)
para("""本文为 X 公司设计的 FSSC 组织架构由"一中心、四模块"组成：财务共享服务中心（设在总部），下设应付管理模块、应收管理模块、费用报销模块、总账核算模块。各模块由 1 名主管 + 2 至 3 名专员组成，总部财务部转型为战略财务，专注于预算、税务、资金等高端业务。""")

heading("业务流程再造", level=2)
para("""以费用报销流程为例，再造前后的对比如下表所示：""")
add_table(
    headers=["环节", "再造前", "再造后", "效率提升"],
    rows=[
        ["凭证审核", "纸质单据 + 人工审核（2 天）", "电子单据 + RPA 自动审核（2 小时）", "85%"],
        ["核算入账", "手工记账（1 天）", "系统自动记账（实时）", "100%"],
        ["费用支付", "线下转账（1 天）", "网银批量支付（10 分钟）", "98%"],
        ["全流程时长", "约 4 个工作日", "约 1 个工作日内", "75%"],
    ],
)

heading("信息系统建设", level=2)
para("""信息系统是 FSSC 运作的技术支撑。本文为 X 公司设计的信息系统采用分层式架构：底层是统一的 ERP 数据库，中间层是 RPA 引擎和影像识别系统，上层是面向员工的自助式报销门户和管理层的财务分析驾驶舱。""")

heading("效益评估", level=2)
para("""FSSC 建设的财务效益可通过成本节约率（Cost Saving Rate, CSR）量化评估，公式如下：""")
# 公式段：用 Typst 数学语法（透传）
para("""$ "CSR" = ("C"_("before") - "C"_("after")) / "C"_("before") times 100% $""", indent=False)

para("""经测算，X 公司 FSSC 建成后预计每年可节约财务成本约 187 万元，成本节约率达 32.7%。同时，效率提升率（Efficiency Improvement Rate, EIR）可达：""")
para("""$ "EIR" = ("T"_("before") - "T"_("after")) / "T"_("before") times 100% = 75% $""", indent=False)


# ============ 第五章 总结与展望 ============
heading("总结与展望", level=1)

heading("研究结论", level=2)
para("""本文以 X 公司为案例对象，研究了数字经济背景下中小企业财务共享服务中心的建设问题，得出以下主要结论：""")
para("""第一，财务共享服务中心是中小企业实现财务数字化转型的有效路径。本文提出的"一中心、四模块"组织架构能够在不大幅增加投入的前提下，实现传统财务管理模式向共享服务模式的平稳过渡。""")
para("""第二，业务流程再造是 FSSC 建设的核心。通过对费用报销等核心流程的重新设计，X 公司预计可以实现 75% 的效率提升。""")
para("""第三，新兴信息技术对 FSSC 的赋能作用显著。RPA、AI、大数据等技术的应用，是中小企业 FSSC 后发优势的关键所在 [12]。""")

heading("研究不足与展望", level=2)
para("""受研究时间和数据可得性的限制，本文还存在以下不足：""")
para("""第一，案例研究的样本数量有限，研究结论的普适性有待进一步验证。未来可以扩展到同行业不同规模的多个案例进行横向比较。""")
para("""第二，对 FSSC 实施过程中的人力资源转型问题探讨不够深入。财务人员从核算型向决策支持型的转变，是 FSSC 长期成功的关键，值得后续研究关注。""")


# ============ 参考文献占位 ============
heading("参考文献", level=1)
doc.add_paragraph("此处由 cnki_refs_accounting.txt 自动接管。")


# ============ 致谢 ============
heading("致谢", level=1)
para("""本论文的完成离不开许多人的帮助和支持，在此我谨向他们表示最诚挚的感谢。""")
para("""首先，我要感谢我的指导老师。从论文选题、文献综述、研究方案设计到论文撰写、修改完善，老师都给予了我无微不至的指导和帮助。老师严谨的治学态度、深厚的专业素养和高尚的师德风范，都将成为我今后学习和工作中的宝贵财富。""")
para("""其次，我要感谢数智经济管理学院会计学专业的各位老师，是你们四年来的悉心教导让我在专业知识和综合素质方面都得到了显著提升。""")
para("""最后，我要感谢我的家人和朋友，是你们一直以来的支持和鼓励让我能够顺利完成学业。由于本人学识有限，论文中难免存在疏漏与不足，恳请各位老师批评指正。""")


doc.save(OUT)
print(f"Sample written: {OUT}")
print(f"Size: {OUT.stat().st_size // 1024} KB")

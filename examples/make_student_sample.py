"""
模拟真实学生工作流：
  - 在 Word 里直接写 [1], [2,3], [1-3] 引用（无需学 BibTeX）
  - 参考文献从知网粘贴出来，独立存为 cnki_refs.txt
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parent / "student_sample.docx"

doc = Document()
style = doc.styles["Normal"]
style.font.name = "宋体"
style.font.size = Pt(12)

# 摘要
p = doc.add_paragraph()
r = p.add_run("摘 要"); r.bold = True; r.font.size = Pt(16)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph(
    "本文研究了基于深度学习的图像分类方法，对比了主流模型架构，"
    "通过实验验证了所提方法的有效性。"
)
p.paragraph_format.first_line_indent = Pt(24)

p = doc.add_paragraph()
r = p.add_run("关键词："); r.bold = True
p.add_run("深度学习；图像分类；卷积神经网络")

# 英文摘要
p = doc.add_paragraph()
r = p.add_run("Abstract"); r.bold = True; r.font.size = Pt(16)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(
    "This thesis studies CNN-based image classification methods."
)
p.paragraph_format.first_line_indent = Pt(24)
p = doc.add_paragraph()
r = p.add_run("Keywords: "); r.bold = True
p.add_run("Deep Learning; Image Classification; CNN")

# 正文 —— 全部用 [N] 中括号引用
doc.add_heading("绪论", level=1)
doc.add_heading("研究背景", level=2)
p = doc.add_paragraph(
    "深度学习已成为计算机视觉领域的核心方法 [4]。"
    "ResNet [5] 提出了残差连接，"
    "Transformer [6] 进一步重塑了模型架构。"
    "中文学者 [1] 在该领域的研究也颇为活跃，"
    "另见 [2,3] 与 [1-3] 等系统性工作。"
)
p.paragraph_format.first_line_indent = Pt(24)

doc.add_heading("研究现状", level=2)
p = doc.add_paragraph(
    "针对图像分类任务，学位论文 [3] 系统总结了现有方法。"
    "近期报告 [7] 指出该领域将持续高速发展。"
)
p.paragraph_format.first_line_indent = Pt(24)

# 第二章
doc.add_heading("方法", level=1)
p = doc.add_paragraph("本文方法基于 ResNet [5] 改进，融合了 Transformer [6] 的注意力机制。")
p.paragraph_format.first_line_indent = Pt(24)

doc.add_heading("实验", level=1)
p = doc.add_paragraph("实验在标准数据集上进行，结果优于基线 [4-6]。")
p.paragraph_format.first_line_indent = Pt(24)

doc.add_heading("总结", level=1)
p = doc.add_paragraph("本文系统研究了图像分类方法。")
p.paragraph_format.first_line_indent = Pt(24)

# 参考文献占位（实际从 cnki_refs.txt 注入）
doc.add_heading("参考文献", level=1)
doc.add_paragraph("此处的内容将由 cnki_refs.txt 自动接管。")

# 致谢
doc.add_heading("致谢", level=1)
p = doc.add_paragraph("感谢导师指导。")
p.paragraph_format.first_line_indent = Pt(24)

doc.save(OUT)
print(f"Sample written: {OUT}")

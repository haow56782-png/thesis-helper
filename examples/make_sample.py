"""
生成一份"乱"的示例 docx：
  - 标题用了 Heading 样式（OK）
  - 但夹杂用 "手动加粗大字号" 当伪标题的段落（lint 应报警）
  - 段落首行没缩进
  - 字号字体多种混用
  - 含一张表格、一段列表、一张图片
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parent / "messy_sample.docx"

doc = Document()

# 标题（用 Heading 1 样式，正确）
doc.add_heading("基于深度学习的图像分类方法研究", level=0)  # title style

# 摘要部分（用大字号手动模拟标题，不规范）
p = doc.add_paragraph()
run = p.add_run("摘 要")
run.bold = True
run.font.size = Pt(16)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    "本文系统地研究了基于卷积神经网络的图像分类方法，重点探讨了数据增强、迁移学习与模型蒸馏三种关键技术对分类精度的影响。"
)

p = doc.add_paragraph()
run = p.add_run("关键词：")
run.bold = True
p.add_run("深度学习；图像分类；卷积神经网络")

# 第一章
doc.add_heading("第一章 绪论", level=1)
doc.add_heading("1.1 研究背景与意义", level=2)
doc.add_paragraph(
    "图像分类是计算机视觉领域最基础也是最重要的任务之一。它的目标是给定一张输入图像，输出该图像所属的类别。"
)
doc.add_paragraph(
    "随着深度学习的崛起，特别是卷积神经网络（CNN）的成功应用，图像分类的精度有了质的飞跃。"
)

doc.add_heading("1.2 本文主要工作", level=2)

# 列表
doc.add_paragraph("复现并对比了主流图像分类模型；", style="List Number")
doc.add_paragraph("提出了一种新的特征融合方法；", style="List Number")
doc.add_paragraph("通过实验验证了方法的有效性。", style="List Number")

# 第二章
doc.add_heading("第二章 方法", level=1)
doc.add_heading("2.1 整体框架", level=2)
doc.add_paragraph(
    "本文提出的方法包含三个主要模块：特征提取、特征融合与分类头。下表展示了核心超参数配置。"
)

# 表格
table = doc.add_table(rows=3, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "模型"
hdr[1].text = "参数量(M)"
hdr[2].text = "Top-1 精度"
table.rows[1].cells[0].text = "ResNet-50"
table.rows[1].cells[1].text = "25.6"
table.rows[1].cells[2].text = "94.2%"
table.rows[2].cells[0].text = "本文方法"
table.rows[2].cells[1].text = "17.9"
table.rows[2].cells[2].text = "94.5%"

# 第三章
doc.add_heading("第三章 实验", level=1)
doc.add_paragraph(
    "实验在 NVIDIA RTX 4090 上进行，数据集采用 CIFAR-10 和 ImageNet-100。结果表明，本文方法在保持精度的前提下减少了 30% 的参数量。"
)

# 一个故意"手动伪标题"
p = doc.add_paragraph()
run = p.add_run("3.1 实验设置（这是手动加粗的伪标题）")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph("详细的实验设置如下：批量大小 256，学习率 0.1，训练 200 轮。")

# 致谢
doc.add_heading("致谢", level=1)
doc.add_paragraph("感谢导师与实验室同学的帮助。")

doc.save(OUT)
print(f"Sample written: {OUT}")

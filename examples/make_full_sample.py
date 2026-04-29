"""
更完整的样本 docx：
  - 完整 5 章结构
  - 摘要+关键词 + 英文摘要+keywords
  - 表格 + 列表 + 二级三级标题
  - 参考文献章节（占位，由 .bib 自动接管）
  - 致谢
  - 引用占位符 @he2016deep 等（直接写在正文，让 Typst 解析）
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parent / "full_sample.docx"

doc = Document()

# 设置默认正文样式（更接近真实学生稿）
style = doc.styles["Normal"]
style.font.name = "宋体"
style.font.size = Pt(12)

# 摘要
p = doc.add_paragraph()
r = p.add_run("摘 要")
r.bold = True; r.font.size = Pt(16)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph(
    "随着深度学习技术的快速发展，图像分类作为计算机视觉领域的核心任务之一，"
    "已经在工业、医疗、安防等多个场景中得到广泛应用。"
    "本文系统地研究了基于卷积神经网络的图像分类方法。"
)
p.paragraph_format.first_line_indent = Pt(24)

p = doc.add_paragraph()
r = p.add_run("关键词：")
r.bold = True
p.add_run("深度学习；图像分类；卷积神经网络；迁移学习")

# 英文摘要
p = doc.add_paragraph()
r = p.add_run("Abstract")
r.bold = True; r.font.size = Pt(16)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph(
    "With the rapid development of deep learning, image classification "
    "has become one of the core tasks in computer vision. "
    "This thesis systematically studies CNN-based image classification methods."
)
p.paragraph_format.first_line_indent = Pt(24)

p = doc.add_paragraph()
r = p.add_run("Keywords: ")
r.bold = True
p.add_run("Deep Learning; Image Classification; CNN; Transfer Learning")

# 第一章
doc.add_heading("绪论", level=1)
doc.add_heading("研究背景与意义", level=2)
p = doc.add_paragraph(
    "图像分类是计算机视觉领域最基础也是最重要的任务之一 @lecun2015deep。"
    "ResNet @he2016deep 提出了残差连接，"
    "Transformer @vaswani2017attention 进一步重塑了模型架构。"
)
p.paragraph_format.first_line_indent = Pt(24)

doc.add_heading("研究现状", level=2)
p = doc.add_paragraph(
    "国内学者 @zhou2016ml 在该领域做出了系统性贡献。"
    "图像分类方法可大致分为以下几类。"
)
p.paragraph_format.first_line_indent = Pt(24)

doc.add_paragraph("基于人工特征的传统方法；", style="List Number")
doc.add_paragraph("基于深度学习的端到端方法；", style="List Number")
doc.add_paragraph("基于自监督学习的预训练方法。", style="List Number")

# 第二章
doc.add_heading("相关工作", level=1)
p = doc.add_paragraph("本章回顾相关工作。详细的对比见下表。")
p.paragraph_format.first_line_indent = Pt(24)

table = doc.add_table(rows=4, cols=4)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "模型"
hdr[1].text = "参数量(M)"
hdr[2].text = "CIFAR-10"
hdr[3].text = "ImageNet"
data = [
    ["ResNet-50", "25.6", "94.2%", "78.3%"],
    ["EfficientNet-B0", "5.3", "93.8%", "77.7%"],
    ["本文方法", "17.9", "94.5%", "78.6%"],
]
for i, row_data in enumerate(data, start=1):
    cells = table.rows[i].cells
    for j, v in enumerate(row_data):
        cells[j].text = v

# 第三章
doc.add_heading("方法设计", level=1)
doc.add_heading("整体框架", level=2)
p = doc.add_paragraph("本文提出的方法包含三个主要模块：特征提取、特征融合与分类头。")
p.paragraph_format.first_line_indent = Pt(24)

doc.add_heading("特征融合策略", level=2)
p = doc.add_paragraph("特征融合的目标是将不同层级的特征有效组合。")
p.paragraph_format.first_line_indent = Pt(24)

# 第四章
doc.add_heading("实验与分析", level=1)
p = doc.add_paragraph(
    "实验在 NVIDIA RTX 4090 上进行，"
    "数据集采用 CIFAR-10 和 ImageNet-100。"
    "结果表明本文方法在保持精度的同时减少了 30% 的参数量。"
)
p.paragraph_format.first_line_indent = Pt(24)

# 第五章
doc.add_heading("总结与展望", level=1)
p = doc.add_paragraph("本文系统研究了基于深度学习的图像分类方法。未来可在更高效的网络结构方向继续优化。")
p.paragraph_format.first_line_indent = Pt(24)

# 参考文献（占位 — 真实数据来自 .bib，这些段落会被 ir_to_typst 跳过）
doc.add_heading("参考文献", level=1)
doc.add_paragraph("[1] 此处的内容会被自动忽略，由 references.bib 自动生成。")

# 致谢
doc.add_heading("致谢", level=1)
p = doc.add_paragraph("感谢导师张教授的悉心指导，感谢实验室同学的帮助与陪伴。")
p.paragraph_format.first_line_indent = Pt(24)

doc.save(OUT)
print(f"Sample written: {OUT}")
